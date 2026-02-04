"""
ExportHandler Lambda Handler

This Lambda function generates exports of document analysis results in multiple formats:
- PDF: Formatted layout with sections
- JSON: Complete data structure
- Excel: Multi-sheet workbook
- Word: Styled document

The exports are uploaded to S3 and a presigned URL is returned for download.

Requirements: 8.1, 8.2, 8.3, 8.4, 8.6, 8.8
"""

import json
import os
import logging
from datetime import datetime
from decimal import Decimal
from typing import Dict, Any, Optional
from io import BytesIO
import tempfile

import boto3
from boto3.dynamodb.conditions import Key
from botocore.exceptions import ClientError

# PDF generation
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib import colors

# Excel generation
from openpyxl import Workbook
from openpyxl.styles import Font, Alignment, PatternFill

# Word generation
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Configure logging
logger = logging.getLogger()
logger.setLevel(logging.INFO)

# Initialize AWS clients
dynamodb = boto3.resource('dynamodb')
s3_client = boto3.client('s3')

# Environment variables
DOCUMENTS_TABLE_NAME = os.environ.get('DOCUMENTS_TABLE_NAME', 'DocumentAnalysis-Documents-dev')
RESULTS_TABLE_NAME = os.environ.get('RESULTS_TABLE_NAME', 'DocumentAnalysis-Results-dev')
RESULTS_BUCKET_NAME = os.environ.get('RESULTS_BUCKET_NAME', 'document-analysis-results-dev')
PRESIGNED_URL_EXPIRATION = int(os.environ.get('PRESIGNED_URL_EXPIRATION', '900'))  # 15 minutes

# Get table references
documents_table = dynamodb.Table(DOCUMENTS_TABLE_NAME)
results_table = dynamodb.Table(RESULTS_TABLE_NAME)


class DecimalEncoder(json.JSONEncoder):
    """Helper class to convert DynamoDB Decimal types to JSON-serializable types"""
    def default(self, obj):
        if isinstance(obj, Decimal):
            return float(obj)
        return super(DecimalEncoder, self).default(obj)


def validate_jwt_token(event: Dict[str, Any]) -> Optional[str]:
    """
    Extract and validate userId from JWT token in API Gateway authorizer context.
    
    Args:
        event: Lambda event from API Gateway
        
    Returns:
        userId if valid, None otherwise
    """
    try:
        authorizer = event.get('requestContext', {}).get('authorizer', {})
        user_id = authorizer.get('claims', {}).get('sub')
        
        if not user_id:
            user_id = authorizer.get('claims', {}).get('cognito:username')
        
        if not user_id:
            logger.error("No userId found in JWT token")
            return None
            
        logger.info(f"Validated userId: {user_id}")
        return user_id
        
    except Exception as e:
        logger.error(f"Error validating JWT token: {str(e)}")
        return None


def get_document_and_analysis(document_id: str, user_id: str) -> Optional[Dict[str, Any]]:
    """
    Retrieve document metadata and analysis results from DynamoDB.
    
    Args:
        document_id: Document ID
        user_id: User ID for authorization
        
    Returns:
        Combined document and analysis data, or None if not found
    """
    try:
        # Get document metadata
        doc_response = documents_table.get_item(Key={'documentId': document_id})
        document = doc_response.get('Item')
        
        if not document:
            logger.error(f"Document not found: {document_id}")
            return None
        
        # Verify user owns the document
        if document.get('userId') != user_id:
            logger.error(f"User {user_id} does not own document {document_id}")
            return None
        
        # Get analysis results
        results_response = results_table.get_item(Key={'documentId': document_id})
        analysis = results_response.get('Item')
        
        if not analysis:
            logger.warning(f"Analysis not found for document: {document_id}")
            analysis = {}
        
        # Combine document and analysis
        combined = {
            'document': document,
            'analysis': analysis
        }
        
        logger.info(f"Retrieved document and analysis for: {document_id}")
        return combined
        
    except ClientError as e:
        logger.error(f"DynamoDB error: {e.response['Error']['Message']}")
        return None
    except Exception as e:
        logger.error(f"Error retrieving document and analysis: {str(e)}")
        return None


def generate_pdf_export(data: Dict[str, Any]) -> bytes:
    """
    Generate PDF export with formatted layout.
    
    Args:
        data: Combined document and analysis data
        
    Returns:
        PDF file as bytes
    """
    try:
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter)
        story = []
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor('#000024'),
            spaceAfter=30,
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=16,
            textColor=colors.HexColor('#008FD0'),
            spaceAfter=12,
        )
        
        # Title
        story.append(Paragraph("Document Analysis Report", title_style))
        story.append(Spacer(1, 0.2 * inch))
        
        # Document metadata
        document = data.get('document', {})
        metadata_data = [
            ['Document Name:', document.get('fileName', 'N/A')],
            ['Vertical:', document.get('vertical', 'N/A').title()],
            ['Upload Date:', document.get('uploadedAt', 'N/A')],
            ['Status:', document.get('status', 'N/A').title()],
            ['File Size:', f"{document.get('fileSize', 0) / 1024:.2f} KB"],
        ]
        
        metadata_table = Table(metadata_data, colWidths=[2*inch, 4*inch])
        metadata_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#F0F0F0')),
            ('TEXTCOLOR', (0, 0), (-1, -1), colors.black),
            ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
            ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (-1, -1), 10),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
            ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ]))
        
        story.append(metadata_table)
        story.append(Spacer(1, 0.3 * inch))
        
        # Analysis results
        analysis = data.get('analysis', {})
        
        if analysis:
            # Executive Summary
            story.append(Paragraph("Executive Summary", heading_style))
            summary_text = analysis.get('executiveSummary', 'No summary available')
            story.append(Paragraph(summary_text, styles['BodyText']))
            story.append(Spacer(1, 0.2 * inch))
            
            # Key Points
            story.append(Paragraph("Key Points", heading_style))
            key_points = analysis.get('keyPoints', [])
            for i, point in enumerate(key_points, 1):
                story.append(Paragraph(f"{i}. {point}", styles['BodyText']))
                story.append(Spacer(1, 0.1 * inch))
            story.append(Spacer(1, 0.2 * inch))
            
            # Next Steps
            story.append(Paragraph("Next Steps", heading_style))
            next_steps = analysis.get('nextSteps', [])
            for i, step in enumerate(next_steps, 1):
                story.append(Paragraph(f"{i}. {step}", styles['BodyText']))
                story.append(Spacer(1, 0.1 * inch))
        else:
            story.append(Paragraph("Analysis not available", styles['BodyText']))
        
        # Footer
        story.append(Spacer(1, 0.5 * inch))
        footer_text = f"Generated on {datetime.utcnow().strftime('%Y-%m-%d %H:%M:%S')} UTC"
        story.append(Paragraph(footer_text, styles['Italic']))
        
        # Build PDF
        doc.build(story)
        
        pdf_bytes = buffer.getvalue()
        buffer.close()
        
        logger.info("Generated PDF export")
        return pdf_bytes
        
    except Exception as e:
        logger.error(f"Error generating PDF: {str(e)}")
        raise


def generate_json_export(data: Dict[str, Any]) -> bytes:
    """
    Generate JSON export with complete data structure.
    
    Args:
        data: Combined document and analysis data
        
    Returns:
        JSON file as bytes
    """
    try:
        # Create export structure
        export_data = {
            'exportedAt': datetime.utcnow().isoformat(),
            'document': data.get('document', {}),
            'analysis': data.get('analysis', {}),
        }
        
        # Convert to JSON with proper formatting
        json_str = json.dumps(export_data, indent=2, cls=DecimalEncoder)
        json_bytes = json_str.encode('utf-8')
        
        logger.info("Generated JSON export")
        return json_bytes
        
    except Exception as e:
        logger.error(f"Error generating JSON: {str(e)}")
        raise


def generate_excel_export(data: Dict[str, Any]) -> bytes:
    """
    Generate Excel export with multiple sheets.
    
    Args:
        data: Combined document and analysis data
        
    Returns:
        Excel file as bytes
    """
    try:
        wb = Workbook()
        
        # Remove default sheet
        wb.remove(wb.active)
        
        # Header styles
        header_font = Font(bold=True, color="FFFFFF")
        header_fill = PatternFill(start_color="000024", end_color="000024", fill_type="solid")
        header_alignment = Alignment(horizontal="left", vertical="center")
        
        # Sheet 1: Document Metadata
        ws_metadata = wb.create_sheet("Document Metadata")
        document = data.get('document', {})
        
        metadata_rows = [
            ['Field', 'Value'],
            ['Document ID', document.get('documentId', 'N/A')],
            ['File Name', document.get('fileName', 'N/A')],
            ['File Type', document.get('fileType', 'N/A')],
            ['File Size (KB)', f"{document.get('fileSize', 0) / 1024:.2f}"],
            ['Vertical', document.get('vertical', 'N/A').title()],
            ['Status', document.get('status', 'N/A').title()],
            ['Uploaded At', document.get('uploadedAt', 'N/A')],
            ['Processed At', document.get('processedAt', 'N/A')],
            ['Processing Time (ms)', document.get('processingTimeMs', 'N/A')],
        ]
        
        for row in metadata_rows:
            ws_metadata.append(row)
        
        # Style header row
        for cell in ws_metadata[1]:
            cell.font = header_font
            cell.fill = header_fill
            cell.alignment = header_alignment
        
        # Adjust column widths
        ws_metadata.column_dimensions['A'].width = 25
        ws_metadata.column_dimensions['B'].width = 50
        
        # Sheet 2: Executive Summary
        analysis = data.get('analysis', {})
        ws_summary = wb.create_sheet("Executive Summary")
        ws_summary.append(['Executive Summary'])
        ws_summary.append([analysis.get('executiveSummary', 'No summary available')])
        
        ws_summary['A1'].font = header_font
        ws_summary['A1'].fill = header_fill
        ws_summary.column_dimensions['A'].width = 100
        ws_summary['A2'].alignment = Alignment(wrap_text=True)
        
        # Sheet 3: Key Points
        ws_keypoints = wb.create_sheet("Key Points")
        ws_keypoints.append(['#', 'Key Point'])
        
        key_points = analysis.get('keyPoints', [])
        for i, point in enumerate(key_points, 1):
            ws_keypoints.append([i, point])
        
        for cell in ws_keypoints[1]:
            cell.font = header_font
            cell.fill = header_fill
        
        ws_keypoints.column_dimensions['A'].width = 5
        ws_keypoints.column_dimensions['B'].width = 100
        
        # Sheet 4: Next Steps
        ws_nextsteps = wb.create_sheet("Next Steps")
        ws_nextsteps.append(['#', 'Next Step'])
        
        next_steps = analysis.get('nextSteps', [])
        for i, step in enumerate(next_steps, 1):
            ws_nextsteps.append([i, step])
        
        for cell in ws_nextsteps[1]:
            cell.font = header_font
            cell.fill = header_fill
        
        ws_nextsteps.column_dimensions['A'].width = 5
        ws_nextsteps.column_dimensions['B'].width = 100
        
        # Save to bytes
        buffer = BytesIO()
        wb.save(buffer)
        excel_bytes = buffer.getvalue()
        buffer.close()
        
        logger.info("Generated Excel export")
        return excel_bytes
        
    except Exception as e:
        logger.error(f"Error generating Excel: {str(e)}")
        raise


def generate_word_export(data: Dict[str, Any]) -> bytes:
    """
    Generate Word export with styled document.
    
    Args:
        data: Combined document and analysis data
        
    Returns:
        Word file as bytes
    """
    try:
        doc = Document()
        
        # Title
        title = doc.add_heading('Document Analysis Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Document Metadata Section
        doc.add_heading('Document Information', level=1)
        
        document = data.get('document', {})
        metadata_table = doc.add_table(rows=6, cols=2)
        metadata_table.style = 'Light Grid Accent 1'
        
        metadata_items = [
            ('Document Name', document.get('fileName', 'N/A')),
            ('Vertical', document.get('vertical', 'N/A').title()),
            ('Upload Date', document.get('uploadedAt', 'N/A')),
            ('Status', document.get('status', 'N/A').title()),
            ('File Size', f"{document.get('fileSize', 0) / 1024:.2f} KB"),
            ('Processing Time', f"{document.get('processingTimeMs', 0)} ms"),
        ]
        
        for i, (label, value) in enumerate(metadata_items):
            row = metadata_table.rows[i]
            row.cells[0].text = label
            row.cells[1].text = str(value)
            # Bold the labels
            row.cells[0].paragraphs[0].runs[0].font.bold = True
        
        doc.add_paragraph()
        
        # Analysis Results
        analysis = data.get('analysis', {})
        
        if analysis:
            # Executive Summary
            doc.add_heading('Executive Summary', level=1)
            summary_text = analysis.get('executiveSummary', 'No summary available')
            doc.add_paragraph(summary_text)
            
            # Key Points
            doc.add_heading('Key Points', level=1)
            key_points = analysis.get('keyPoints', [])
            for point in key_points:
                doc.add_paragraph(point, style='List Bullet')
            
            # Next Steps
            doc.add_heading('Next Steps', level=1)
            next_steps = analysis.get('nextSteps', [])
            for i, step in enumerate(next_steps, 1):
                doc.add_paragraph(f"{i}. {step}", style='List Number')
        else:
            doc.add_paragraph('Analysis not available')
        
        # Footer
        doc.add_paragraph()
        footer_para = doc.add_paragraph()
        footer_run = footer_para.add_run(
            f"Generated on {datetime.utcnow().strftime('%Y-%m-%d %H:%M:%S')} UTC"
        )
        footer_run.font.italic = True
        footer_run.font.size = Pt(9)
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Save to bytes
        buffer = BytesIO()
        doc.save(buffer)
        word_bytes = buffer.getvalue()
        buffer.close()
        
        logger.info("Generated Word export")
        return word_bytes
        
    except Exception as e:
        logger.error(f"Error generating Word: {str(e)}")
        raise


def upload_to_s3_and_get_url(file_bytes: bytes, document_id: str, user_id: str, 
                              file_format: str) -> str:
    """
    Upload export file to S3 and generate presigned URL.
    
    Args:
        file_bytes: File content as bytes
        document_id: Document ID
        user_id: User ID
        file_format: Export format (pdf, json, excel, word)
        
    Returns:
        Presigned URL for download
    """
    try:
        # Determine file extension and content type
        format_config = {
            'pdf': {'ext': 'pdf', 'content_type': 'application/pdf'},
            'json': {'ext': 'json', 'content_type': 'application/json'},
            'excel': {'ext': 'xlsx', 'content_type': 'application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'},
            'word': {'ext': 'docx', 'content_type': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'},
        }
        
        config = format_config.get(file_format, {'ext': 'bin', 'content_type': 'application/octet-stream'})
        
        # S3 key
        s3_key = f"exports/{user_id}/{document_id}/export.{config['ext']}"
        
        # Upload to S3
        s3_client.put_object(
            Bucket=RESULTS_BUCKET_NAME,
            Key=s3_key,
            Body=file_bytes,
            ContentType=config['content_type'],
        )
        
        logger.info(f"Uploaded export to S3: {s3_key}")
        
        # Generate presigned URL
        presigned_url = s3_client.generate_presigned_url(
            'get_object',
            Params={
                'Bucket': RESULTS_BUCKET_NAME,
                'Key': s3_key,
            },
            ExpiresIn=PRESIGNED_URL_EXPIRATION
        )
        
        logger.info(f"Generated presigned URL with {PRESIGNED_URL_EXPIRATION}s expiration")
        return presigned_url
        
    except ClientError as e:
        logger.error(f"S3 error: {e.response['Error']['Message']}")
        raise
    except Exception as e:
        logger.error(f"Error uploading to S3: {str(e)}")
        raise


def lambda_handler(event: Dict[str, Any], context: Any) -> Dict[str, Any]:
    """
    Main Lambda handler for ExportHandler.
    
    Args:
        event: API Gateway event
        context: Lambda context
        
    Returns:
        API Gateway response with download URL
    """
    try:
        logger.info(f"Received event: {json.dumps(event)}")
        
        # Validate JWT token and extract userId
        user_id = validate_jwt_token(event)
        if not user_id:
            return {
                'statusCode': 401,
                'headers': {
                    'Content-Type': 'application/json',
                    'Access-Control-Allow-Origin': '*',
                },
                'body': json.dumps({
                    'error': 'Unauthorized',
                    'message': 'Invalid or missing authentication token'
                })
            }
        
        # Extract document ID from path parameters
        document_id = event.get('pathParameters', {}).get('documentId')
        if not document_id:
            return {
                'statusCode': 400,
                'headers': {
                    'Content-Type': 'application/json',
                    'Access-Control-Allow-Origin': '*',
                },
                'body': json.dumps({
                    'error': 'Bad Request',
                    'message': 'Missing documentId in path'
                })
            }
        
        # Extract format from request body
        body = json.loads(event.get('body', '{}'))
        export_format = body.get('format', 'pdf').lower()
        
        if export_format not in ['pdf', 'json', 'excel', 'word']:
            return {
                'statusCode': 400,
                'headers': {
                    'Content-Type': 'application/json',
                    'Access-Control-Allow-Origin': '*',
                },
                'body': json.dumps({
                    'error': 'Bad Request',
                    'message': 'Invalid format. Must be one of: pdf, json, excel, word'
                })
            }
        
        # Retrieve document and analysis
        data = get_document_and_analysis(document_id, user_id)
        if not data:
            return {
                'statusCode': 404,
                'headers': {
                    'Content-Type': 'application/json',
                    'Access-Control-Allow-Origin': '*',
                },
                'body': json.dumps({
                    'error': 'Not Found',
                    'message': 'Document not found or access denied'
                })
            }
        
        # Generate export based on format
        if export_format == 'pdf':
            file_bytes = generate_pdf_export(data)
        elif export_format == 'json':
            file_bytes = generate_json_export(data)
        elif export_format == 'excel':
            file_bytes = generate_excel_export(data)
        elif export_format == 'word':
            file_bytes = generate_word_export(data)
        
        # Upload to S3 and get presigned URL
        download_url = upload_to_s3_and_get_url(file_bytes, document_id, user_id, export_format)
        
        # Return response
        return {
            'statusCode': 200,
            'headers': {
                'Content-Type': 'application/json',
                'Access-Control-Allow-Origin': '*',
            },
            'body': json.dumps({
                'downloadUrl': download_url,
                'expiresIn': PRESIGNED_URL_EXPIRATION,
                'format': export_format,
                'documentId': document_id,
                'generatedAt': datetime.utcnow().isoformat()
            })
        }
        
    except Exception as e:
        logger.error(f"Error in lambda_handler: {str(e)}", exc_info=True)
        return {
            'statusCode': 500,
            'headers': {
                'Content-Type': 'application/json',
                'Access-Control-Allow-Origin': '*',
            },
            'body': json.dumps({
                'error': 'Internal Server Error',
                'message': 'Failed to generate export'
            })
        }
