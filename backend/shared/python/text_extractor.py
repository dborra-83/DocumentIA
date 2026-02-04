"""
Text extraction utilities for document processing.

This module provides functions to extract text from various document formats:
- PDF files using PyPDF2
- DOCX files using python-docx
- TXT files using standard file reading

All functions include error handling for corrupted or invalid files.
"""

import io
import logging
from typing import Optional

try:
    # Try to import pypdf (newer, recommended)
    from pypdf import PdfReader
    from pypdf.errors import PdfReadError
except ImportError:
    # Fallback to PyPDF2 if pypdf is not available
    from PyPDF2 import PdfReader
    from PyPDF2.errors import PdfReadError

from docx import Document

logger = logging.getLogger(__name__)


class TextExtractionError(Exception):
    """Exception raised when text extraction fails."""
    pass


def extract_text_from_pdf(file_content: bytes) -> str:
    """
    Extract text from a PDF file.
    
    Args:
        file_content: Raw bytes of the PDF file
        
    Returns:
        Extracted text as a string
        
    Raises:
        TextExtractionError: If the PDF is corrupted or cannot be read
    """
    try:
        pdf_file = io.BytesIO(file_content)
        pdf_reader = PdfReader(pdf_file)
        
        # Check if PDF is encrypted
        if pdf_reader.is_encrypted:
            logger.warning("PDF is encrypted, attempting to decrypt")
            try:
                pdf_reader.decrypt('')
            except Exception as e:
                raise TextExtractionError(f"Failed to decrypt PDF: {str(e)}")
        
        # Extract text from all pages
        text_parts = []
        num_pages = len(pdf_reader.pages)
        
        if num_pages == 0:
            raise TextExtractionError("PDF has no pages")
        
        logger.info(f"Extracting text from {num_pages} pages")
        
        for page_num in range(num_pages):
            try:
                page = pdf_reader.pages[page_num]
                text = page.extract_text()
                if text:
                    text_parts.append(text)
            except Exception as e:
                logger.warning(f"Failed to extract text from page {page_num + 1}: {str(e)}")
                continue
        
        extracted_text = '\n\n'.join(text_parts)
        
        if not extracted_text.strip():
            raise TextExtractionError("No text could be extracted from PDF")
        
        logger.info(f"Successfully extracted {len(extracted_text)} characters from PDF")
        return extracted_text
        
    except PdfReadError as e:
        raise TextExtractionError(f"Invalid or corrupted PDF file: {str(e)}")
    except Exception as e:
        if isinstance(e, TextExtractionError):
            raise
        raise TextExtractionError(f"Failed to extract text from PDF: {str(e)}")


def extract_text_from_docx(file_content: bytes) -> str:
    """
    Extract text from a DOCX file.
    
    Args:
        file_content: Raw bytes of the DOCX file
        
    Returns:
        Extracted text as a string
        
    Raises:
        TextExtractionError: If the DOCX is corrupted or cannot be read
    """
    try:
        docx_file = io.BytesIO(file_content)
        doc = Document(docx_file)
        
        # Extract text from all paragraphs
        text_parts = []
        
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text:
                text_parts.append(text)
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        text_parts.append(text)
        
        extracted_text = '\n\n'.join(text_parts)
        
        if not extracted_text.strip():
            raise TextExtractionError("No text could be extracted from DOCX")
        
        logger.info(f"Successfully extracted {len(extracted_text)} characters from DOCX")
        return extracted_text
        
    except Exception as e:
        if isinstance(e, TextExtractionError):
            raise
        raise TextExtractionError(f"Failed to extract text from DOCX: {str(e)}")


def extract_text_from_txt(file_content: bytes) -> str:
    """
    Extract text from a TXT file.
    
    Args:
        file_content: Raw bytes of the TXT file
        
    Returns:
        Extracted text as a string
        
    Raises:
        TextExtractionError: If the TXT file cannot be decoded
    """
    try:
        # Try UTF-8 first
        try:
            text = file_content.decode('utf-8')
        except UnicodeDecodeError:
            # Fallback to latin-1 which accepts all byte values
            logger.warning("UTF-8 decoding failed, trying latin-1")
            text = file_content.decode('latin-1')
        
        if not text.strip():
            raise TextExtractionError("TXT file is empty")
        
        logger.info(f"Successfully extracted {len(text)} characters from TXT")
        return text
        
    except Exception as e:
        if isinstance(e, TextExtractionError):
            raise
        raise TextExtractionError(f"Failed to extract text from TXT: {str(e)}")


def extract_text(file_content: bytes, file_type: str) -> str:
    """
    Extract text from a document based on its file type.
    
    This is a convenience function that routes to the appropriate
    extraction function based on file type.
    
    Args:
        file_content: Raw bytes of the file
        file_type: File type ('pdf', 'docx', or 'txt')
        
    Returns:
        Extracted text as a string
        
    Raises:
        TextExtractionError: If extraction fails or file type is unsupported
        ValueError: If file_type is not supported
    """
    file_type = file_type.lower().strip()
    
    if file_type == 'pdf':
        return extract_text_from_pdf(file_content)
    elif file_type in ['docx', 'doc']:
        return extract_text_from_docx(file_content)
    elif file_type == 'txt':
        return extract_text_from_txt(file_content)
    else:
        raise ValueError(f"Unsupported file type: {file_type}")
