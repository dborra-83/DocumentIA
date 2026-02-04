"""
Unit tests for text extraction utilities.

Tests cover:
- PDF text extraction
- DOCX text extraction
- TXT text extraction
- Error handling for corrupted files
- Edge cases (empty files, encrypted PDFs, etc.)
"""

import io
import pytest
from unittest.mock import Mock, patch, MagicMock

try:
    from pypdf import PdfWriter
    from pypdf.errors import PdfReadError
except ImportError:
    from PyPDF2 import PdfWriter
    from PyPDF2.errors import PdfReadError

from docx import Document
from docx.text.paragraph import Paragraph

from text_extractor import (
    extract_text_from_pdf,
    extract_text_from_docx,
    extract_text_from_txt,
    extract_text,
    TextExtractionError
)


class TestExtractTextFromPDF:
    """Tests for PDF text extraction."""
    
    def test_extract_text_from_valid_pdf(self):
        """Test extracting text from a valid PDF with content."""
        # Create a simple PDF in memory
        pdf_buffer = io.BytesIO()
        pdf_writer = PdfWriter()
        
        # Add a page with text
        page = pdf_writer.add_blank_page(width=200, height=200)
        pdf_writer.write(pdf_buffer)
        pdf_content = pdf_buffer.getvalue()
        
        # Mock the PDF reader to return test text
        with patch('text_extractor.PdfReader') as mock_reader:
            mock_page = Mock()
            mock_page.extract_text.return_value = "This is test content from a PDF."
            
            mock_pdf = Mock()
            mock_pdf.is_encrypted = False
            mock_pdf.pages = [mock_page]
            mock_reader.return_value = mock_pdf
            
            result = extract_text_from_pdf(pdf_content)
            
            assert result == "This is test content from a PDF."
            assert len(result) > 0
    
    def test_extract_text_from_multi_page_pdf(self):
        """Test extracting text from a PDF with multiple pages."""
        pdf_content = b"fake pdf content"
        
        with patch('text_extractor.PdfReader') as mock_reader:
            mock_page1 = Mock()
            mock_page1.extract_text.return_value = "Page 1 content"
            
            mock_page2 = Mock()
            mock_page2.extract_text.return_value = "Page 2 content"
            
            mock_pdf = Mock()
            mock_pdf.is_encrypted = False
            mock_pdf.pages = [mock_page1, mock_page2]
            mock_reader.return_value = mock_pdf
            
            result = extract_text_from_pdf(pdf_content)
            
            assert "Page 1 content" in result
            assert "Page 2 content" in result
            assert result == "Page 1 content\n\nPage 2 content"
    
    def test_extract_text_from_encrypted_pdf(self):
        """Test handling encrypted PDF files."""
        pdf_content = b"fake encrypted pdf"
        
        with patch('text_extractor.PdfReader') as mock_reader:
            mock_page = Mock()
            mock_page.extract_text.return_value = "Decrypted content"
            
            mock_pdf = Mock()
            mock_pdf.is_encrypted = True
            mock_pdf.decrypt = Mock()
            mock_pdf.pages = [mock_page]
            mock_reader.return_value = mock_pdf
            
            result = extract_text_from_pdf(pdf_content)
            
            mock_pdf.decrypt.assert_called_once_with('')
            assert result == "Decrypted content"
    
    def test_extract_text_from_encrypted_pdf_decrypt_fails(self):
        """Test handling encrypted PDF that cannot be decrypted."""
        pdf_content = b"fake encrypted pdf"
        
        with patch('text_extractor.PdfReader') as mock_reader:
            mock_pdf = Mock()
            mock_pdf.is_encrypted = True
            mock_pdf.decrypt.side_effect = Exception("Cannot decrypt")
            mock_reader.return_value = mock_pdf
            
            with pytest.raises(TextExtractionError) as exc_info:
                extract_text_from_pdf(pdf_content)
            
            assert "Failed to decrypt PDF" in str(exc_info.value)
    
    def test_extract_text_from_empty_pdf(self):
        """Test handling PDF with no pages."""
        pdf_content = b"fake pdf"
        
        with patch('text_extractor.PdfReader') as mock_reader:
            mock_pdf = Mock()
            mock_pdf.is_encrypted = False
            mock_pdf.pages = []
            mock_reader.return_value = mock_pdf
            
            with pytest.raises(TextExtractionError) as exc_info:
                extract_text_from_pdf(pdf_content)
            
            assert "PDF has no pages" in str(exc_info.value)
    
    def test_extract_text_from_pdf_no_extractable_text(self):
        """Test handling PDF with pages but no extractable text."""
        pdf_content = b"fake pdf"
        
        with patch('text_extractor.PdfReader') as mock_reader:
            mock_page = Mock()
            mock_page.extract_text.return_value = ""
            
            mock_pdf = Mock()
            mock_pdf.is_encrypted = False
            mock_pdf.pages = [mock_page]
            mock_reader.return_value = mock_pdf
            
            with pytest.raises(TextExtractionError) as exc_info:
                extract_text_from_pdf(pdf_content)
            
            assert "No text could be extracted" in str(exc_info.value)
    
    def test_extract_text_from_corrupted_pdf(self):
        """Test handling corrupted PDF files."""
        pdf_content = b"not a valid pdf"
        
        with patch('text_extractor.PdfReader') as mock_reader:
            mock_reader.side_effect = PdfReadError("Invalid PDF")
            
            with pytest.raises(TextExtractionError) as exc_info:
                extract_text_from_pdf(pdf_content)
            
            assert "Invalid or corrupted PDF" in str(exc_info.value)
    
    def test_extract_text_from_pdf_page_extraction_fails(self):
        """Test handling when some pages fail to extract."""
        pdf_content = b"fake pdf"
        
        with patch('text_extractor.PdfReader') as mock_reader:
            mock_page1 = Mock()
            mock_page1.extract_text.return_value = "Page 1 content"
            
            mock_page2 = Mock()
            mock_page2.extract_text.side_effect = Exception("Page error")
            
            mock_page3 = Mock()
            mock_page3.extract_text.return_value = "Page 3 content"
            
            mock_pdf = Mock()
            mock_pdf.is_encrypted = False
            mock_pdf.pages = [mock_page1, mock_page2, mock_page3]
            mock_reader.return_value = mock_pdf
            
            result = extract_text_from_pdf(pdf_content)
            
            # Should still extract from successful pages
            assert "Page 1 content" in result
            assert "Page 3 content" in result


class TestExtractTextFromDOCX:
    """Tests for DOCX text extraction."""
    
    def test_extract_text_from_valid_docx(self):
        """Test extracting text from a valid DOCX file."""
        docx_content = b"fake docx content"
        
        with patch('text_extractor.Document') as mock_document:
            mock_para1 = Mock()
            mock_para1.text = "First paragraph"
            
            mock_para2 = Mock()
            mock_para2.text = "Second paragraph"
            
            mock_doc = Mock()
            mock_doc.paragraphs = [mock_para1, mock_para2]
            mock_doc.tables = []
            mock_document.return_value = mock_doc
            
            result = extract_text_from_docx(docx_content)
            
            assert "First paragraph" in result
            assert "Second paragraph" in result
            assert result == "First paragraph\n\nSecond paragraph"
    
    def test_extract_text_from_docx_with_tables(self):
        """Test extracting text from DOCX with tables."""
        docx_content = b"fake docx content"
        
        with patch('text_extractor.Document') as mock_document:
            mock_para = Mock()
            mock_para.text = "Document text"
            
            # Mock table
            mock_cell = Mock()
            mock_cell.text = "Table cell content"
            
            mock_row = Mock()
            mock_row.cells = [mock_cell]
            
            mock_table = Mock()
            mock_table.rows = [mock_row]
            
            mock_doc = Mock()
            mock_doc.paragraphs = [mock_para]
            mock_doc.tables = [mock_table]
            mock_document.return_value = mock_doc
            
            result = extract_text_from_docx(docx_content)
            
            assert "Document text" in result
            assert "Table cell content" in result
    
    def test_extract_text_from_docx_with_empty_paragraphs(self):
        """Test extracting text from DOCX with empty paragraphs."""
        docx_content = b"fake docx content"
        
        with patch('text_extractor.Document') as mock_document:
            mock_para1 = Mock()
            mock_para1.text = "Content"
            
            mock_para2 = Mock()
            mock_para2.text = "   "  # Whitespace only
            
            mock_para3 = Mock()
            mock_para3.text = ""  # Empty
            
            mock_doc = Mock()
            mock_doc.paragraphs = [mock_para1, mock_para2, mock_para3]
            mock_doc.tables = []
            mock_document.return_value = mock_doc
            
            result = extract_text_from_docx(docx_content)
            
            # Should only include non-empty paragraphs
            assert result == "Content"
    
    def test_extract_text_from_empty_docx(self):
        """Test handling empty DOCX files."""
        docx_content = b"fake docx content"
        
        with patch('text_extractor.Document') as mock_document:
            mock_doc = Mock()
            mock_doc.paragraphs = []
            mock_doc.tables = []
            mock_document.return_value = mock_doc
            
            with pytest.raises(TextExtractionError) as exc_info:
                extract_text_from_docx(docx_content)
            
            assert "No text could be extracted" in str(exc_info.value)
    
    def test_extract_text_from_corrupted_docx(self):
        """Test handling corrupted DOCX files."""
        docx_content = b"not a valid docx"
        
        with patch('text_extractor.Document') as mock_document:
            mock_document.side_effect = Exception("Invalid DOCX")
            
            with pytest.raises(TextExtractionError) as exc_info:
                extract_text_from_docx(docx_content)
            
            assert "Failed to extract text from DOCX" in str(exc_info.value)


class TestExtractTextFromTXT:
    """Tests for TXT text extraction."""
    
    def test_extract_text_from_valid_txt_utf8(self):
        """Test extracting text from a valid UTF-8 TXT file."""
        text_content = "This is a test text file.\nWith multiple lines."
        txt_content = text_content.encode('utf-8')
        
        result = extract_text_from_txt(txt_content)
        
        assert result == text_content
        assert "test text file" in result
    
    def test_extract_text_from_txt_with_special_characters(self):
        """Test extracting text with special characters."""
        text_content = "Special chars: é, ñ, ü, 中文"
        txt_content = text_content.encode('utf-8')
        
        result = extract_text_from_txt(txt_content)
        
        assert result == text_content
    
    def test_extract_text_from_txt_latin1_fallback(self):
        """Test fallback to latin-1 encoding when UTF-8 fails."""
        # Create content that's valid latin-1 but not UTF-8
        txt_content = b'\xe9\xe0\xfc'  # Valid latin-1, invalid UTF-8
        
        result = extract_text_from_txt(txt_content)
        
        # Should successfully decode with latin-1
        assert len(result) > 0
    
    def test_extract_text_from_empty_txt(self):
        """Test handling empty TXT files."""
        txt_content = b""
        
        with pytest.raises(TextExtractionError) as exc_info:
            extract_text_from_txt(txt_content)
        
        assert "TXT file is empty" in str(exc_info.value)
    
    def test_extract_text_from_whitespace_only_txt(self):
        """Test handling TXT files with only whitespace."""
        txt_content = b"   \n\n   \t  "
        
        with pytest.raises(TextExtractionError) as exc_info:
            extract_text_from_txt(txt_content)
        
        assert "TXT file is empty" in str(exc_info.value)


class TestExtractText:
    """Tests for the main extract_text function."""
    
    def test_extract_text_routes_to_pdf(self):
        """Test that extract_text routes PDF files correctly."""
        pdf_content = b"fake pdf"
        
        with patch('text_extractor.extract_text_from_pdf') as mock_pdf:
            mock_pdf.return_value = "PDF content"
            
            result = extract_text(pdf_content, 'pdf')
            
            mock_pdf.assert_called_once_with(pdf_content)
            assert result == "PDF content"
    
    def test_extract_text_routes_to_docx(self):
        """Test that extract_text routes DOCX files correctly."""
        docx_content = b"fake docx"
        
        with patch('text_extractor.extract_text_from_docx') as mock_docx:
            mock_docx.return_value = "DOCX content"
            
            result = extract_text(docx_content, 'docx')
            
            mock_docx.assert_called_once_with(docx_content)
            assert result == "DOCX content"
    
    def test_extract_text_routes_to_txt(self):
        """Test that extract_text routes TXT files correctly."""
        txt_content = b"fake txt"
        
        with patch('text_extractor.extract_text_from_txt') as mock_txt:
            mock_txt.return_value = "TXT content"
            
            result = extract_text(txt_content, 'txt')
            
            mock_txt.assert_called_once_with(txt_content)
            assert result == "TXT content"
    
    def test_extract_text_handles_case_insensitive_type(self):
        """Test that file type is case-insensitive."""
        txt_content = b"fake txt"
        
        with patch('text_extractor.extract_text_from_txt') as mock_txt:
            mock_txt.return_value = "TXT content"
            
            result = extract_text(txt_content, 'TXT')
            
            mock_txt.assert_called_once()
            assert result == "TXT content"
    
    def test_extract_text_handles_whitespace_in_type(self):
        """Test that file type whitespace is handled."""
        txt_content = b"fake txt"
        
        with patch('text_extractor.extract_text_from_txt') as mock_txt:
            mock_txt.return_value = "TXT content"
            
            result = extract_text(txt_content, '  txt  ')
            
            mock_txt.assert_called_once()
            assert result == "TXT content"
    
    def test_extract_text_unsupported_file_type(self):
        """Test handling unsupported file types."""
        content = b"fake content"
        
        with pytest.raises(ValueError) as exc_info:
            extract_text(content, 'xlsx')
        
        assert "Unsupported file type" in str(exc_info.value)
    
    def test_extract_text_doc_alias(self):
        """Test that 'doc' is treated as 'docx'."""
        doc_content = b"fake doc"
        
        with patch('text_extractor.extract_text_from_docx') as mock_docx:
            mock_docx.return_value = "DOC content"
            
            result = extract_text(doc_content, 'doc')
            
            mock_docx.assert_called_once_with(doc_content)
            assert result == "DOC content"
